import streamlit as st
import tempfile, os, shutil, re, json, time, gc
from collections import Counter
from bs4 import BeautifulSoup
from docx import Document as DocxDocument

from langchain_community.llms import Ollama
from langchain_community.embeddings import OllamaEmbeddings
from langchain_community.document_loaders import PyPDFLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.docstore.document import Document
from langchain.vectorstores import Chroma

st.set_page_config(layout="wide", page_title="NotebookLM + Chat - Fixed Timed Reset")

st.title("📒 NotebookLM + Chat (Fixed Timed Reset)")

embedder = OllamaEmbeddings(model="nomic-embed-text:latest")
llm = Ollama(model="llama3.1:8b-instruct-q5_K_M")

with st.sidebar:
    st.header("Knowledge Base Projects")
    kb_root_dir = "kb_projects"
    os.makedirs(kb_root_dir, exist_ok=True)
    kb_list = [d for d in os.listdir(kb_root_dir) if os.path.isdir(os.path.join(kb_root_dir, d))]

    project = st.selectbox("Select KB Project", ["Create new..."] + kb_list)

    if project == "Create new...":
        new_name = st.text_input("New KB name")
        if st.button("Create KB") and new_name:
            os.makedirs(os.path.join(kb_root_dir, new_name), exist_ok=True)
            st.success(f"Created: {new_name}")
            st.rerun()

    else:
        kb_dir = os.path.join(kb_root_dir, project)
        del_flag = os.path.join(kb_dir, "DELETE_ON_RESTART.flag")
        sel_file_path = os.path.join(kb_dir, "selected_sources.json")

        # Check on startup: delete if flagged
        if os.path.exists(del_flag):
            if os.path.exists(kb_dir):
                try:
                    shutil.rmtree(kb_dir, ignore_errors=True)
                except:
                    pass
                os.makedirs(kb_dir, exist_ok=True)
            try:
                os.remove(del_flag)
            except:
                pass
            st.success("Old knowledge base removed. Fresh workspace created.")

        # Reset KB button
        if st.button("🚨 Reset Current KB"):
            try:
                if 'vectorstore' in st.session_state and st.session_state.vectorstore:
                    st.session_state.vectorstore._client.reset()
            except:
                pass
            st.session_state.vectorstore = None
            gc.collect()

            with open(del_flag, "w") as f:
                f.write("DELETE THIS KB ON RESTART\n")

            st.warning("The app will close in 5 seconds to reset this KB. Please restart it manually after.")
            st.toast("App will close in 5 seconds...", icon="⏰")
            time.sleep(5)
            st.toast("Exiting now. Please restart app.", icon="⚠️")
            os._exit(0)

        top_k = st.slider("Chunks to retrieve (k)", 1, 20, 10)
        show_refs = st.checkbox("Show references in answers", value=True)

        uploaded_files = st.file_uploader("Browse files", accept_multiple_files=True, type=["pdf", "txt", "docx", "html"])

        if 'vectorstore' not in st.session_state:
            try:
                st.session_state.vectorstore = Chroma(persist_directory=kb_dir, embedding_function=embedder)
                st.session_state.docs_uploaded = True
            except:
                st.session_state.vectorstore = None
                st.session_state.docs_uploaded = False

        docs_list = []
        if st.session_state.vectorstore:
            docs_list = st.session_state.vectorstore.similarity_search("", k=1000)
        files = sorted(set([d.metadata.get('source', 'Unknown') for d in docs_list]))

        if os.path.exists(sel_file_path):
            with open(sel_file_path, "r") as f:
                saved_selection = set(json.load(f))
        else:
            saved_selection = set(files)
            with open(sel_file_path, "w") as f:
                json.dump(list(saved_selection), f)

        select_all = st.checkbox("Select all sources", value=len(saved_selection) == len(files))
        new_selection = set()

        for fname in files:
            checked = fname in saved_selection
            if st.checkbox(fname, checked, key=f"file_{fname}"):
                new_selection.add(fname)

        if select_all:
            new_selection = set(files)

        with open(sel_file_path, "w") as f:
            json.dump(list(new_selection), f)

        st.session_state.selected_files = new_selection

        if st.button("Add to Knowledge Base") and uploaded_files:
            docs = []
            for file in uploaded_files:
                suffix = os.path.splitext(file.name)[1].lower()

                if suffix == ".pdf":
                    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_file:
                        tmp_file.write(file.read())
                        tmp_file.flush()
                        tmp_path = tmp_file.name
                    loader = PyPDFLoader(tmp_path)
                    pdf_docs = loader.load()
                    os.remove(tmp_path)
                    full_text = " ".join([d.page_content for d in pdf_docs])
                elif suffix == ".docx":
                    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_file:
                        tmp_file.write(file.read())
                        tmp_file.flush()
                        tmp_path = tmp_file.name
                    docx_obj = DocxDocument(tmp_path)
                    full_text = "\n".join([para.text for para in docx_obj.paragraphs])
                    os.remove(tmp_path)
                elif suffix == ".html":
                    html_content = file.read().decode("utf-8")
                    soup = BeautifulSoup(html_content, "html.parser")
                    full_text = soup.get_text(separator="\n")
                else:
                    full_text = file.read().decode("utf-8")

                words = re.findall(r'\w+', full_text.lower())
                common = Counter(words).most_common(20)
                tags = ", ".join([w for w, _ in common[:5]])

                metadata = {"source": file.name, "tags": tags}

                splitter = RecursiveCharacterTextSplitter(chunk_size=800, chunk_overlap=200)
                split_texts = splitter.split_text(full_text)
                for chunk_text in split_texts:
                    docs.append(Document(page_content=chunk_text, metadata=metadata))

            prev_docs = []
            if st.session_state.vectorstore:
                prev_docs = st.session_state.vectorstore.similarity_search("", k=1000)
            all_docs = prev_docs + docs
            texts = [d.page_content for d in all_docs]
            metadatas = [d.metadata for d in all_docs]

            st.session_state.vectorstore = Chroma.from_texts(
                texts, embedder, metadatas=metadatas, persist_directory=kb_dir
            )
            st.session_state.vectorstore.persist()
            st.success("Files added!")
            st.rerun()

# Main panel
if project != "Create new...":

    if 'chat_history' not in st.session_state:
        st.session_state.chat_history = []

    if 'sources' not in st.session_state:
        st.session_state.sources = {}

    selected_files = st.session_state.selected_files
    count_sel = len(selected_files)
    st.write(f"### {count_sel} sources selected")

    if count_sel > 0:
        docs = st.session_state.vectorstore.similarity_search("", k=1000)
        selected_docs = [d for d in docs if d.metadata.get('source') in selected_files]
        combined_text = ""

        for d in selected_docs:
            combined_text += d.page_content[:1000] + "\n"
            if len(combined_text) > 4000:
                combined_text = combined_text[:4000]
                break

        words = re.findall(r'\w+', combined_text.lower())
        common = Counter(words).most_common(30)
        keywords = [w for w, _ in common][:10]

        prompt_sum = f"Provide a brief, combined summary of these sources with key points:\n{combined_text}"
        summary = llm(prompt_sum)

        for kw in keywords:
            summary = re.sub(rf"(?i)\b({re.escape(kw)})\b", r"**\1**", summary)

        st.write("#### Overview")
        st.write(summary)
    else:
        st.info("No sources selected — chat will use no document context.")

    user_q = st.text_input("Ask a question (always available):")

    if st.button("Reset Chat"):
        st.session_state.chat_history = []

    if user_q:
        st.session_state.chat_history.append(("user", user_q))
        retrieval_query = f"Retrieve info for: {user_q}"

        retrieved = []
        if count_sel > 0:
            results = st.session_state.vectorstore.similarity_search(retrieval_query, k=top_k)
            for doc in results:
                if doc.metadata.get('source') in selected_files:
                    retrieved.append(doc)

        context = ""
        sources = {}
        for idx, doc in enumerate(retrieved):
            snippet = doc.page_content
            for word in user_q.split():
                snippet = re.sub(rf"(?i)\b({re.escape(word)})\b", r"**\1**", snippet)
            context += f"[{idx+1}] {snippet}\n"
            sources[idx+1] = (doc.metadata.get("source", "Unknown"), snippet)

        if show_refs:
            prompt = f"""Based only on the following context, answer with citations like [n]. If no context, answer generally.

Context:
{context}

Question: {user_q}

Answer:"""
        else:
            prompt = f"""Based only on the following context, answer WITHOUT citations. If no context, answer generally.

Context:
{context}

Question: {user_q}

Answer:"""

        answer = llm(prompt)
        st.session_state.chat_history.append(("assistant", answer))
        st.session_state.sources = sources

    for who, msg in st.session_state.chat_history[::-1]:
        if who == "user":
            st.markdown(f"**You:** {msg}")
        else:
            if show_refs:
                def citation_link(match):
                    num = int(match.group(1))
                    return f'<a href="#source{num}">[{num}]</a>'
                msg_html = re.sub(r"\[(\d+)\]", citation_link, msg)
                st.markdown(f"**Assistant:** {msg_html}", unsafe_allow_html=True)
            else:
                st.markdown(f"**Assistant:** {msg}")

    if show_refs and st.session_state.sources:
        st.subheader("Sources")
        for idx, (filename, snippet) in st.session_state.sources.items():
            st.markdown(f'<a name="source{idx}"></a>', unsafe_allow_html=True)
            with st.expander(f"Source [{idx}] - {filename}"):
                st.write(snippet)
